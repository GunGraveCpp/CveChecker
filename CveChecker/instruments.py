import functools
import os
from sys import stderr

import requests
import vulners
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm

from CveChecker.config import *


def time_measurement(func):
    '''Декоратор для замера времени'''

    import time

    @functools.wraps(func)
    def inner(*args, **kwargs):
        start = time.time()
        ret_value = func(*args, **kwargs)
        end = time.time()
        print('Function {} worked: {} min {} sec'.format(func.__name__,
                                                         int(end - start) // 60, int(end - start) % 60))
        return ret_value

    return inner if DEBUG else func


def trace(func):
    '''Декоратор для отладки'''

    @functools.wraps(func)
    def inner(*args, **kwargs):
        ret_value = func(*args, **kwargs)
        print('Function {} \n\tInput args: {} \n\tInput kwargs: {} \n\tReturn {}'.format(
            func.__name__, args, kwargs, ret_value), end='\n\n')
        return ret_value

    return inner if DEBUG else func


def final_message(text):
    '''Декоратор для вывода сообщения'''

    def wrapper(func):
        @functools.wraps(func)
        def inner(*args, **kwargs):
            ret_value = func(*args, **kwargs)
            print(text)
            return ret_value

        return inner if INFO else func

    return wrapper


def read_package_in_file(name):
    package_list = {'deb': [], 'rpm': [], 'archive': [], 'unknown': []}

    with open(name, 'r') as file:
        for package in file:
            if package.rfind('.deb') != -1:
                package_list['deb'].append(
                    package.replace('_', ' ').replace('.deb', ''))
            elif package.rfind('.rpm') != -1:
                package_list['rpm'].append(package.replace('.rpm', ''))
            # package.count('-') костыль, чтобы была возможность делить имя и версию архива
            elif (package.rfind('.tar') & package.rfind('.zip') != -1) & package.count('_'):
                package_list['archive'].append(
                    package[:package.find('.tar') & package.find('.zip')])
            else:
                package_list['unknown'].append(package)

    return package_list


def set_col_widths(table):
    widths = (Cm(width) for width in COLUMN_SIZE)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def generate_doc_table(document):
    document.add_heading(TABLE_NAME, 1)
    table = document.add_table(rows=1, cols=6, style='Table Grid')

    for column in range(0, len(table.columns)):
        table.cell(0, column).text = TABLE_HEADER[column]

    return table


def deb_version(name):
    return name.split()[1]


def rpm_version(name):
    return name.rsplit('.', 1)[0].rsplit('-', 1)[0].rsplit('-', 1)[1]


def archive_version(name):
    return name.split('_', 1)[1]


def get_description(cve):
    '''  Функция получения описпния CVE
    Получение описания и наличия eploit с https://nvd.nist.gov/
    по идентификатору cve
    '''

    req = requests.get(NIST_LINK.format(cve))
    answ = {'description': WITHOUT_DESCRIPTION,
            'exploit': AVAILABILITY_EPLOIT_NO,
            'cvss': WITHOUT_DESCRIPTION}

    if req.ok:
        soup = BeautifulSoup(req.text, 'lxml')

        span_list = soup.find_all('span', class_="badge")
        if span_list:
            for span in span_list:
                if span.text == 'Exploit':
                    answ['exploit'] = AVAILABILITY_EPLOIT_YES
                    break
        cvss = soup.find('span', {'data-testid': 'vuln-cvss2-panel-vector'})
        if cvss:
            answ['cvss'] = cvss.text[1:-1]  # Обрезаем скобки у вектора cvss

        desc = soup.find('p', {'data-testid': 'vuln-description'})
        if desc:
            answ['description'] = desc.text

    return answ


def fill_package_info(package_info, cve_list):
    '''Заполнение информации о cve'''

    for cve in cve_list:
        cve_info = get_description(cve)

        package_info['cvelist'].append(
            {'cve': cve,
             'description': cve_info['description'],
             'exploit': cve_info['exploit'],
             'cvss': cve_info['cvss']
             }
        )


def get_package_info_deb(name, vulners_api):
    '''Получение информации о deb пакете'''

    package_info = {'package': name.replace(
        ' ', '_'), 'version': deb_version(name), 'cvelist': []}
    try:
        results = vulners_api.os_audit(
            os=DEB_OS, version=DEB_VERSION, packages=[name])

        if len(results['reasons']) != 0:
            cve_list = results.get('cvelist')
            fill_package_info(package_info, cve_list)

    except Exception as er:
        print(er, file=stderr)

    return package_info


def get_package_info_rpm(name, vulners_api):
    '''Получение информации о rpm пакете'''

    package_info = {'package': name,
                    'version': rpm_version(name), 'cvelist': []}
    try:
        results = vulners_api.os_audit(
            os=CENT_OS, version=CENT_VERSION, packages=[name])

        if len(results['reasons']) != 0:
            cve_list = results.get('cvelist')
            fill_package_info(package_info, cve_list)

    except Exception as er:
        print(er, file=stderr)

    return package_info


def get_package_info_archive(name, vulners_api):
    '''Получение информации об архиве'''

    package_info = {'package': name.split(
        '_', 1)[0], 'version': archive_version(name), 'cvelist': []}

    try:
        results = vulners_api.get_software_vulnerabilities(
            package_info['package'].upper(), package_info['version'])

        cve_list = list()

        if len(results) != 0:
            # Valners рандомно может выдать request  с разными корнями
            for cve in results['NVD' if "NVD" in results else "software"]:
                if cve['cvelist'] is not None:
                    cve_list += cve['cvelist']
            fill_package_info(package_info, cve_list)

    except Exception as er:
        print(er, file=stderr)

    return package_info


def past_package_in_table(table, package_info):
    '''Вставки информации о пакете в таблицу'''

    try:
        for cve_info in package_info['cvelist']:
            row = table.add_row()
            row.cells[0].text = cve_info['cve']
            row.cells[1].text = cve_info['cvss']
            row.cells[2].text = cve_info['description']
            row.cells[3].text = package_info['package'] + " " + package_info['version']
            row.cells[5].text = cve_info['exploit']

    except Exception as er:
        print(er, file=stderr)


def gather_information(table, package_list):
    '''Сбор информации о пакетах'''

    vulners_api = vulners.VulnersApi(api_key=os.environ["VULNERS_API_KEY"])
    package_list
    for package_deb in package_list['deb']:
        package_info = get_package_info_deb(package_deb, vulners_api)
        past_package_in_table(table, package_info)
    for package_rpm in package_list['rpm']:
        package_info = get_package_info_rpm(package_rpm, vulners_api)
        past_package_in_table(table, package_info)
    for package_archive in package_list['archive']:
        package_info = get_package_info_archive(package_archive, vulners_api)
        past_package_in_table(table, package_info)
    with open("../unknown.txt", "w") as outfile:
        outfile.write('\n'.join(package_list['unknown']))


@time_measurement
def run_check(inputfile, outputfile):
    '''Запуск задачи проверки и генерация отчета'''
    try:
        document = Document()
        package_list = read_package_in_file(inputfile)
        table = generate_doc_table(document)
        gather_information(table, package_list)
        set_col_widths(table)
        document.save(outputfile)
    except Exception as er:
        print(er, file=stderr)
